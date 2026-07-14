import csv
import io
import json
import re
import time
import uuid
from pathlib import Path
from typing import Optional

from fastapi import HTTPException, UploadFile

from utils.supabase_client import supabase

UPLOAD_DIR = Path("uploads")
ALLOWED_TYPES = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".csv": "text/csv",
    ".json": "application/json",
}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_TEXT_LENGTH = 10 * 1024 * 1024  # 10 MB equivalent for pasted text
MISSING_TABLE_MSG = "Database table knowledge_files is missing. Please run schema_new_tables.sql in Supabase SQL Editor."


def _validate_bot_ownership(bot_id: str, user_id: str, workspace_id: Optional[str] = None):
    try:
        bot = supabase.table("bots").select("id, clerk_user_id, workspace_id").eq("id", bot_id).execute()
        if not bot.data:
            raise HTTPException(status_code=404, detail="Bot not found")
        b = bot.data[0]
        if b.get("clerk_user_id") != user_id:
            try:
                member = supabase.table("bot_members").select("*").eq("bot_id", bot_id).eq("clerk_user_id", user_id).execute()
                if not member.data or not member.data[0].get("member_email", "").endswith("[edit]"):
                    raise HTTPException(status_code=403, detail="Access denied: bot does not belong to you or you lack edit access.")
            except HTTPException:
                raise
            except Exception:
                raise HTTPException(status_code=403, detail="Access denied: bot does not belong to you")
        return b
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid bot_id: {e}")


def _is_missing_table_error(e: Exception) -> bool:
    msg = str(e)
    return "PGRST205" in msg or "Could not find the table" in msg or ("relation" in msg and "does not exist" in msg)


def _set_status(file_id: str, status: str, error: Optional[str] = None):
    update = {"status": status}
    if error is not None:
        update["error_message"] = error
    try:
        supabase.table("knowledge_files").update(update).eq("id", file_id).execute()
    except Exception:
        pass


def _extract_text_fallback(file_path: Path) -> str:
    try:
        return file_path.read_text(encoding="utf-8", errors="replace").strip()
    except Exception:
        return ""


def _strip_markdown(text: str) -> str:
    """Strip markdown syntax while preserving section boundaries."""
    # Preserve heading text but mark boundaries
    text = re.sub(r'^#{1,6}\s+(.+)$', r'\n\1\n', text, flags=re.MULTILINE)
    # Remove bold/italic markers
    text = re.sub(r'\*{1,3}(.+?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}(.+?)_{1,3}', r'\1', text)
    # Convert links [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Remove images ![alt](url)
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    # Remove inline code backticks
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove code block fences
    text = re.sub(r'^```[\s\S]*?```', '', text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r'^>\s?', '', text, flags=re.MULTILINE)
    # Remove horizontal rules
    text = re.sub(r'^[-*_]{3,}\s*$', '\n', text, flags=re.MULTILINE)
    # Remove list markers but keep text
    text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
    # Clean up excessive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _extract_csv(file_path: Path) -> str:
    """Parse CSV into readable text rows."""
    try:
        raw = file_path.read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(io.StringIO(raw))
        rows = list(reader)
        if not rows:
            return ""
        headers = rows[0] if rows else []
        lines = []
        for row in rows[1:]:
            parts = [f"{headers[i]}: {val}" if i < len(headers) else val for i, val in enumerate(row) if val.strip()]
            if parts:
                lines.append("; ".join(parts))
        return "\n".join(lines).strip()
    except Exception:
        return file_path.read_text(encoding="utf-8", errors="replace").strip()


def _extract_json(file_path: Path) -> str:
    """Extract readable text values from JSON."""
    try:
        raw = file_path.read_text(encoding="utf-8", errors="replace")
        data = json.loads(raw)

        def _collect_strings(obj, depth=0):
            parts = []
            if isinstance(obj, str):
                if obj.strip():
                    parts.append(obj.strip())
            elif isinstance(obj, dict):
                for k, v in obj.items():
                    nested = _collect_strings(v, depth + 1)
                    if nested:
                        parts.append(f"{k}: {nested}")
            elif isinstance(obj, list):
                for item in obj:
                    nested = _collect_strings(item, depth + 1)
                    if nested:
                        parts.append(nested)
            return "\n".join(parts) if depth == 0 else "; ".join(parts)

        return _collect_strings(data)
    except Exception:
        return file_path.read_text(encoding="utf-8", errors="replace").strip()


def _extract_text(file_path: Path, ext: str) -> str:
    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(str(file_path))
            pages = []
            for i, page in enumerate(doc):
                page_text = page.get_text().strip()
                if page_text:
                    pages.append(f"\n\n--- Page {i + 1} ---\n\n{page_text}")
            doc.close()
            return "\n".join(pages).strip()
        except ImportError:
            text = _extract_text_fallback(file_path)
            if text:
                return text
            raise HTTPException(status_code=400, detail="No readable text found in this PDF. Install PyMuPDF (pip install PyMuPDF) for better PDF support.")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {e}")
    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(str(file_path))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Preserve heading structure
                    if para.style and para.style.name and para.style.name.startswith('Heading'):
                        parts.append(f"\n{para.text.strip()}\n")
                    else:
                        parts.append(para.text.strip())
            # Extract tables
            for table in doc.tables:
                headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_parts = [
                        f"{headers[i]}: {val}" if i < len(headers) and headers[i] else val
                        for i, val in enumerate(cells) if val
                    ]
                    if row_parts:
                        parts.append("; ".join(row_parts))
            return "\n".join(parts).strip()
        except ImportError:
            text = _extract_text_fallback(file_path)
            if text:
                return text
            raise HTTPException(status_code=400, detail="No readable text found in this DOCX. Install python-docx (pip install python-docx) for better DOCX support.")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {e}")
    elif ext == ".md":
        try:
            raw = file_path.read_text(encoding="utf-8", errors="replace").strip()
            return _strip_markdown(raw)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read markdown file: {e}")
    elif ext == ".csv":
        text = _extract_csv(file_path)
        if not text:
            raise HTTPException(status_code=400, detail="CSV file is empty or could not be parsed.")
        return text
    elif ext == ".json":
        text = _extract_json(file_path)
        if not text:
            raise HTTPException(status_code=400, detail="JSON file is empty or could not be parsed.")
        return text
    else:
        try:
            return file_path.read_text(encoding="utf-8", errors="replace").strip()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read file: {e}")


def _process_file_text(file_id: str, bot_id: Optional[str], text: str):
    from services.chunking import chunk_faq_text
    from services.embedding import embed_texts

    # Remove null bytes to prevent PostgreSQL errors
    text = text.replace("\x00", "").replace("\u0000", "")

    chunks = chunk_faq_text(text) if len(text) > 20 else [text]
    if not chunks:
        chunks = [text]

    embeddings = embed_texts(chunks)

    documents = [
        {"bot_id": bot_id, "chunk_text": chunk, "embedding": emb}
        for chunk, emb in zip(chunks, embeddings)
    ] if bot_id else []

    if documents:
        for i in range(0, len(documents), 3):
            batch = documents[i:i + 3]
            for attempt in range(1, 4):
                try:
                    supabase.table("documents").insert(batch).execute()
                    break
                except Exception:
                    if attempt == 3:
                        raise
                    time.sleep(attempt)

    supabase.table("knowledge_files").update({
        "status": "embedded",
        "chunks_count": len(chunks),
    }).eq("id", file_id).execute()


def list_files(user_id: str, bot_id: str):
    try:
        query = supabase.table("knowledge_files").select("*").eq("bot_id", bot_id).order("created_at", desc=True)
        result = query.execute()
        return result.data
    except Exception as e:
        if _is_missing_table_error(e):
            return []
        return []


def get_file(file_id: str, user_id: str, bot_id: Optional[str] = None):
    try:
        query = supabase.table("knowledge_files").select("*").eq("id", file_id)
        if bot_id:
            query = query.eq("bot_id", bot_id)
        result = query.execute()
        if not result.data:
            raise HTTPException(status_code=404, detail="File not found")
        return result.data[0]
    except HTTPException:
        raise
    except Exception as e:
        if _is_missing_table_error(e):
            raise HTTPException(status_code=500, detail=MISSING_TABLE_MSG)
        raise HTTPException(status_code=404, detail="File not found")


async def upload_file(file: UploadFile, user_id: str, uploaded_by: str, bot_id: str):
    if not bot_id:
        raise HTTPException(status_code=400, detail="bot_id is required to upload a file")

    ext = Path(file.filename).suffix.lower() if file.filename else ""
    if ext not in ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail=f"File type {ext} not supported. Allowed: {', '.join(ALLOWED_TYPES.keys())}")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File exceeds 10 MB limit")

    file_id = str(uuid.uuid4())
    safe_name = f"{file_id}{ext}"
    file_path = UPLOAD_DIR / safe_name
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

    with open(file_path, "wb") as f:
        f.write(content)

    row = {
        "id": file_id,
        "bot_id": bot_id,
        "user_id": user_id,
        "file_name": safe_name,
        "original_name": file.filename,
        "file_type": ext.lstrip(".").upper(),
        "file_size": len(content),
        "storage_path": str(file_path),
        "status": "pending",
    }

    try:
        result = supabase.table("knowledge_files").insert(row).execute()
        record = result.data[0]
    except Exception as e:
        if file_path.exists():
            file_path.unlink()
        if _is_missing_table_error(e):
            raise HTTPException(status_code=500, detail=MISSING_TABLE_MSG)
        with open("upload_error.log", "a") as f:
            f.write(f"Upload Error: {e}\n")
        raise HTTPException(status_code=500, detail="Upload failed: " + str(e))

    _set_status(file_id, "processing")
    try:
        text = _extract_text(file_path, ext)
        if not text:
            _set_status(file_id, "failed", "No text could be extracted from this file.")
            try:
                result = supabase.table("knowledge_files").select("*").eq("id", file_id).execute()
                if result.data:
                    record = result.data[0]
            except Exception:
                pass
            return record

        _process_file_text(file_id, bot_id, text)
    except HTTPException:
        _set_status(file_id, "failed", "Text extraction failed")
        raise
    except Exception as e:
        with open("upload_error.log", "a") as f:
            f.write(f"Processing Error: {e}\n")
        _set_status(file_id, "failed", str(e))
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

    try:
        result = supabase.table("knowledge_files").select("*").eq("id", file_id).execute()
        if result.data:
            record = result.data[0]
    except Exception:
        pass
    return record


def delete_file(file_id: str, user_id: str, bot_id: str):
    record = get_file(file_id, user_id, bot_id)
    file_path = Path(record["storage_path"])
    if file_path.exists():
        file_path.unlink()
    try:
        supabase.table("knowledge_files").delete().eq("id", file_id).eq("bot_id", bot_id).execute()
        supabase.table("documents").delete().eq("bot_id", bot_id).execute()
    except Exception:
        pass
    return True


def reprocess_file(file_id: str, user_id: str, bot_id: str):
    record = get_file(file_id, user_id, bot_id)
    file_path = Path(record["storage_path"])
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    ext = Path(record["file_name"]).suffix.lower()
    _set_status(file_id, "processing")

    try:
        text = _extract_text(file_path, ext)
        if not text:
            _set_status(file_id, "failed", "No text could be extracted.")
            return

        _process_file_text(file_id, bot_id, text)
    except HTTPException:
        _set_status(file_id, "failed", "Processing failed")
        raise
    except Exception as e:
        _set_status(file_id, "failed", str(e))
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")


def export_knowledge(user_id: str, bot_id: str):
    try:
        files = list_files(user_id, bot_id)
    except Exception:
        files = []
    export_data = []
    for f in files:
        file_path = Path(f["storage_path"])
        content = ""
        if file_path.exists():
            try:
                ext = Path(f["file_name"]).suffix.lower()
                content = _extract_text(file_path, ext)
            except Exception:
                content = "[Binary content - not readable as text]"
        export_data.append({
            "file_name": f["original_name"],
            "file_type": f["file_type"],
            "file_size": f["file_size"],
            "embedding_status": f["status"],
            "uploaded_at": f["created_at"],
            "content": content[:10000],
        })
    return {"files": export_data, "total": len(export_data)}


def upload_text(title: str, content: str, user_id: str, bot_id: str):
    """Upload pasted text directly as knowledge."""
    if not bot_id:
        raise HTTPException(status_code=400, detail="bot_id is required")
    if not content or not content.strip():
        raise HTTPException(status_code=400, detail="Content cannot be empty")
    if len(content) > MAX_TEXT_LENGTH:
        raise HTTPException(status_code=400, detail="Text exceeds 10 MB limit")

    # Sanitize title
    safe_title = re.sub(r'[^\w\s\-.]', '', title.strip())[:100] or "Pasted Text"
    if not safe_title.endswith('.txt'):
        safe_title += '.txt'

    file_id = str(uuid.uuid4())
    file_path = UPLOAD_DIR / f"{file_id}.txt"
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

    content_bytes = content.encode("utf-8")
    with open(file_path, "wb") as f:
        f.write(content_bytes)

    row = {
        "id": file_id,
        "bot_id": bot_id,
        "user_id": user_id,
        "file_name": f"{file_id}.txt",
        "original_name": safe_title,
        "file_type": "TXT",
        "file_size": len(content_bytes),
        "storage_path": str(file_path),
        "status": "pending",
    }

    try:
        result = supabase.table("knowledge_files").insert(row).execute()
        record = result.data[0]
    except Exception as e:
        if file_path.exists():
            file_path.unlink()
        if _is_missing_table_error(e):
            raise HTTPException(status_code=500, detail=MISSING_TABLE_MSG)
        raise HTTPException(status_code=500, detail="Upload failed: " + str(e))

    _set_status(file_id, "processing")
    try:
        # Remove null bytes
        clean_text = content.replace("\x00", "").replace("\u0000", "").strip()
        if not clean_text:
            _set_status(file_id, "failed", "No text content found.")
            return record

        _process_file_text(file_id, bot_id, clean_text)
    except HTTPException:
        _set_status(file_id, "failed", "Processing failed")
        raise
    except Exception as e:
        _set_status(file_id, "failed", str(e))
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

    try:
        result = supabase.table("knowledge_files").select("*").eq("id", file_id).execute()
        if result.data:
            record = result.data[0]
    except Exception:
        pass
    return record
